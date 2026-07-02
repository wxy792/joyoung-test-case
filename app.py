#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
九阳程序用例转换网站后端
支持两种转换：
1. 菜单流程 -> 菜单一致性核对用例
2. 报警文件 -> 报警用例
"""

import os
import io
import json
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import FormulaRule
from flask import Flask, render_template, request, send_file, flash, redirect, url_for
import shutil, copy, re, glob
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__,
            template_folder=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates'),
            static_folder=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static'))
app.secret_key = 'joyoung-test-case-tool-2026'
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # 禁用静态文件缓存
app.config['TEMPLATES_AUTO_RELOAD'] = True  # 模板自动重载

UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
OUTPUT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'outputs')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


# ===================== 通用样式 =====================

def apply_header_style(cell):
    """表头样式：深蓝底白字、加粗、居中"""
    cell.font = Font(name='微软雅黑', bold=True, color='FFFFFF', size=11)
    cell.fill = PatternFill(start_color='2F5496', end_color='2F5496', fill_type='solid')
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin = Side(style='thin', color='B8CCE4')
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def apply_data_style(cell, is_step_col=False):
    """数据行样式"""
    cell.font = Font(name='微软雅黑', size=10)
    cell.alignment = Alignment(horizontal='center' if is_step_col else 'left',
                                vertical='center', wrap_text=True)
    thin = Side(style='thin', color='B8CCE4')
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)


def auto_width(sheet, min_width=8, max_width=60):
    """自动调整列宽"""
    for col in sheet.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                val_len = len(str(cell.value))
                # 考虑中文字符
                chinese_count = sum(1 for c in str(cell.value) if '\u4e00' <= c <= '\u9fff')
                effective_len = val_len + chinese_count
                max_len = max(max_len, min(effective_len, max_width))
        adjusted = min(max(max_len + 2, min_width), max_width)
        sheet.column_dimensions[col_letter].width = adjusted


# ===================== 转换功能1：菜单一致性核对 =====================

def convert_menu_consistency(input_path, output_path):
    """
    将菜单流程图转换为菜单一致性核对用例
    处理逻辑：
    1. 遍历所有sheet（跳过名为'目录'的sheet）
    2. 删除D列（实时时间/倒计时）
    3. 在D列填入'一致性评估'表头，各行填入'符合:□     不符合:□'
    4. 检查A列步骤序号是否连续，空单元格也补充数字
    5. 所有单元格水平居中、自动换行
    6. 自动调整行高，避免文字被覆盖
    """
    wb = openpyxl.load_workbook(input_path)
    new_wb = openpyxl.Workbook()
    new_wb.remove(new_wb.active)

    for sheet_name in wb.sheetnames:
        if sheet_name == '目录':
            continue

        src_sheet = wb[sheet_name]
        new_sheet = new_wb.create_sheet(title=sheet_name)

        all_rows = list(src_sheet.iter_rows(values_only=True))
        if not all_rows:
            continue

        # 找到表头行
        header_row_idx = None
        for i, row in enumerate(all_rows):
            row_vals = [str(v).strip() if v else '' for v in row]
            joined = ''.join(row_vals)
            if '步骤' in joined and ('动作' in joined or '标准时间' in joined):
                header_row_idx = i
                break

        if header_row_idx is None:
            header_row_idx = 1

        new_row_idx = 1
        expected_step = 1

        # 处理标题行（表头之前的行）
        for i in range(header_row_idx):
            row_vals = all_rows[i]
            if all(v is None or str(v).strip() == '' for v in row_vals):
                continue
            title_val = next((v for v in row_vals if v is not None and str(v).strip()), '')
            cell = new_sheet.cell(row=new_row_idx, column=1, value=title_val)
            cell.font = Font(name='微软雅黑', bold=True, size=14, color='1e40af')
            cell.alignment = Alignment(horizontal='center', vertical='center')
            new_sheet.merge_cells(f'A{new_row_idx}:E{new_row_idx}')
            new_sheet.row_dimensions[new_row_idx].height = 28
            new_row_idx += 1

        # 处理表头行
        new_header = ['步骤', '动作', '标准时间', '一致性评估', '备注']
        for col_idx, val in enumerate(new_header, 1):
            cell = new_sheet.cell(row=new_row_idx, column=col_idx, value=val)
            apply_header_style(cell)
        new_sheet.row_dimensions[new_row_idx].height = 30
        header_output_row = new_row_idx
        new_row_idx += 1

        # 处理数据行
        for row_idx in range(header_row_idx + 1, len(all_rows)):
            row_vals = list(all_rows[row_idx])

            # 跳过完全空的行
            if all(v is None or str(v).strip() == '' for v in row_vals):
                continue

            # 原数据列: A=步骤, B=动作, C=标准时间, D=实时时间(丢弃), E=备注
            step_val = row_vals[0] if len(row_vals) > 0 else None
            action_val = row_vals[1] if len(row_vals) > 1 else None
            std_time_val = row_vals[2] if len(row_vals) > 2 else None
            remark_val = row_vals[4] if len(row_vals) > 4 else None

            # 需求1：不管A列是否为空，都补充连续的步骤序号
            step_num = expected_step
            expected_step += 1
            new_sheet.cell(row=new_row_idx, column=1, value=step_num)

            # 需求2：B列（动作）为空时，A列原文字放入B列，B、C合并
            # 适用于：原文件某些行A列有说明文字、B/C列为空的情况
            merge_bc = False
            if action_val is None or str(action_val).strip() == '':
                merge_bc = True
                # A列原文字作为B列内容（B、C合并显示），D列留空（下拉选择）
                header_text = str(step_val).strip() if (step_val is not None and str(step_val).strip()) else ''
                new_sheet.cell(row=new_row_idx, column=2, value=header_text)
            else:
                # 正常数据行，D列留空（下拉选择）
                new_sheet.cell(row=new_row_idx, column=2, value=action_val)
                new_sheet.cell(row=new_row_idx, column=3, value=std_time_val)
                new_sheet.cell(row=new_row_idx, column=5, value=remark_val)

            # 需求3+4：所有单元格水平居中、自动换行、自适应行高
            max_lines = 1
            cols_to_style = [1, 2, 4, 5] if merge_bc else [1, 2, 3, 4, 5]
            for col_idx in cols_to_style:
                cell = new_sheet.cell(row=new_row_idx, column=col_idx)
                cell.font = Font(name='微软雅黑', size=10)
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                thin = Side(style='thin', color='B8CCE4')
                cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
                if cell.value:
                    lines = str(cell.value).count('\n') + 1
                    max_lines = max(max_lines, lines)

            if merge_bc:
                new_sheet.merge_cells(f'B{new_row_idx}:C{new_row_idx}')
                # 合并后重新设定B列对齐方式
                b_cell = new_sheet.cell(row=new_row_idx, column=2)
                b_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

            new_sheet.row_dimensions[new_row_idx].height = max(20, max_lines * 18)

            new_row_idx += 1

        # 为D列（一致性评估）添加数据验证：下拉选择「符合/不符合」，二选一
        # 方案：在隐藏列（F列）写入选项，然后引用该列作为下拉源（兼容性最好）
        if new_row_idx > header_output_row + 1:
            # 在F列（第6列）写入选项（隐藏该列）
            new_sheet.cell(row=1, column=6, value='符合')
            new_sheet.cell(row=2, column=6, value='不符合')
            new_sheet.column_dimensions['F'].hidden = True

            # 创建数据验证，引用F1:F2作为选项源
            dv = DataValidation(
                type="list",
                formula1="=$F$1:$F$2",
                allow_blank=True,
                prompt="请选择一致性评估结果",
                promptTitle="一致性评估"
            )
            col_d_range = f'D{header_output_row + 1}:D{new_row_idx - 1}'
            dv.add(col_d_range)
            new_sheet.add_data_validation(dv)

        # 为D列添加条件格式：值为「不符合」时红底强调
        if new_row_idx > header_output_row + 1:
            red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
            col_d_range = f'D{header_output_row + 1}:D{new_row_idx - 1}'
            # FormulaRule: 用公式 D3="不符合" 匹配（条件格式会自动按行调整）
            rule = FormulaRule(formula=[f'D{header_output_row + 1}="不符合"'], fill=red_fill)
            new_sheet.conditional_formatting.add(col_d_range, rule)

        # 设置列宽
        col_widths = [8, 40, 12, 20, 30]
        for i, width in enumerate(col_widths, 1):
            new_sheet.column_dimensions[get_column_letter(i)].width = width

        # 冻结首行
        new_sheet.freeze_panes = f'A{header_output_row + 1}'

    new_wb.save(output_path)
    return len(new_wb.sheetnames)


# ===================== 报警核对Word生成（完整版） =====================

def set_cell_text_keep_valign(cell, text, bold=False, size=10.5, align='center'):
    """设置单元格文字，保留vAlign XML属性（垂直居中）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign_elem = tcPr.find(qn('w:vAlign'))
    vAlign_val = vAlign_elem.get(qn('w:val')) if vAlign_elem is not None else None

    cell.text = ''
    para = cell.paragraphs[0]
    if align == 'left':
        para.alignment = 0
    else:
        para.alignment = 1

    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if vAlign_val:
        new_vAlign = OxmlElement('w:vAlign')
        new_vAlign.set(qn('w:val'), vAlign_val)
        old = tcPr.find(qn('w:vAlign'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(new_vAlign)
    cell.vertical_alignment = 1


def parse_trigger_states(raw_states):
    """健壮解析触发状态（兼容多种格式）"""
    if not raw_states or not raw_states.strip():
        return ['']
    parts = re.split(r'(?=\s*\d+\.)', raw_states)
    parts = [p.strip() for p in parts if p.strip()]
    if len(parts) > 1:
        cleaned = [re.sub(r'^\s*\d+\.\s*', '', p) for p in parts]
        return cleaned
    lines = [l.strip() for l in raw_states.split('\n') if l.strip()]
    if len(lines) > 1:
        return lines
    return [raw_states.strip()]


def parse_alarm_code(alarm_type_str):
    """解析报警代码和名称（支持中英文括号）"""
    s = alarm_type_str.strip()
    match = re.match(r'([A-Za-z0-9]+)\s*[（(]([^）)]+)[）)]', s)
    if match:
        return match.group(1), match.group(2)
    parts = re.split(r'\s+', s, 1)
    if len(parts) >= 2:
        return parts[0], parts[1]
    return s, s


def find_alarm_table(doc):
    """动态查找报警表格（不依赖固定索引）"""
    for ti, table in enumerate(doc.tables):
        try:
            first_row = table.rows[0]
            headers = [cell.text.strip() for cell in first_row.cells]
            header_str = ' '.join(headers)
            if '报警类型' in header_str or '触发状态' in header_str:
                return table
        except:
            continue
    if doc.tables:
        return doc.tables[-1]
    return None


def extract_alarms_from_spec(spec_path):
    """从线路板功能说明书提取报警列表（动态列映射）"""
    doc = Document(spec_path)
    alarm_table = find_alarm_table(doc)
    if alarm_table is None:
        return []

    headers = [cell.text.strip() for cell in alarm_table.rows[0].cells]
    col_map = {}
    for hi, h in enumerate(headers):
        h = h.strip()
        if '序号' in h and '序号' not in col_map:
            col_map['序号'] = hi
        elif '报警类型' in h and '报警类型' not in col_map:
            col_map['报警类型'] = hi
        elif '报警文本' in h and '报警文本' not in col_map:
            col_map['报警文本'] = hi
        elif '触发状态' in h and '触发状态' not in col_map:
            col_map['触发状态'] = hi
        elif '触发功能' in h and '触发功能' not in col_map:
            col_map['触发功能'] = hi
        elif '报警原因' in h or '触发方法' in h:
            col_map['报警原因及触发方法'] = hi
        elif '是否可恢复' in h and '是否可恢复' not in col_map:
            col_map['是否可恢复'] = hi

    defaults = {'序号': 0, '报警类型': 1, '报警文本': 2, '触发状态': 3,
                '触发功能': 4, '报警原因及触发方法': 5, '是否可恢复': 6}
    for k, v in defaults.items():
        if k not in col_map:
            col_map[k] = v

    alarms = []
    for i, row in enumerate(alarm_table.rows):
        if i == 0:
            continue
        cells = [cell.text.strip() for cell in row.cells]
        while len(cells) <= max(col_map.values()):
            cells.append('')
        alarms.append({
            '序号': cells[col_map['序号']],
            '报警类型': cells[col_map['报警类型']],
            '报警文本': cells[col_map['报警文本']],
            '触发状态': cells[col_map['触发状态']],
            '触发功能': cells[col_map['触发功能']],
            '报警原因及触发方法': cells[col_map['报警原因及触发方法']],
            '是否可恢复': cells[col_map['是否可恢复']]
        })
    return alarms


def generate_full_alarm_report(template_path, spec_path, output_path, title_info=None):
    """
    完整报警核对报告生成：
    - 复制模板文档，清除原内容后插入新表格
    - 根据触发状态数量自动复制报警表格
    - 填充所有预期结果列
    """
    if title_info is None:
        title_info = {}

    alarms = extract_alarms_from_spec(spec_path)
    if len(alarms) == 0:
        return 0

    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))

    for child in list(body):
        if child.tag != (qn('w:sectPr')) and child is not sectPr:
            body.remove(child)

    tpl = Document(template_path)
    tpl_title_tbl = tpl.tables[0]
    new_title_tbl = copy.deepcopy(tpl_title_tbl._tbl)
    body.insert(0, new_title_tbl)

    title_tbl = doc.tables[0]
    set_cell_text_keep_valign(title_tbl.rows[1].cells[1], title_info.get('model', ''), size=9)
    set_cell_text_keep_valign(title_tbl.rows[2].cells[1], title_info.get('voltage', ''), size=9)

    power_cell = title_tbl.rows[2].cells[3]
    power_lines = title_info.get('power', ['', ''])
    tc = power_cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign_elem = tcPr.find(qn('w:vAlign'))
    vAlign_val = vAlign_elem.get(qn('w:val')) if vAlign_elem is not None else None
    power_cell.text = ''
    para = power_cell.paragraphs[0]
    para.alignment = 1
    for j, line in enumerate(power_lines):
        run = para.add_run(line)
        run.font.size = Pt(9)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if j < len(power_lines) - 1:
            br = OxmlElement('w:br')
            run._element.append(br)
    if vAlign_val:
        new_va = OxmlElement('w:vAlign')
        new_va.set(qn('w:val'), vAlign_val)
        old = tcPr.find(qn('w:vAlign'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(new_va)
    power_cell.vertical_alignment = 1

    total_tables = 0
    if len(tpl.tables) > 1:
        tpl_alarm_tbl = tpl.tables[1]

        for idx, alarm in enumerate(alarms):
            code, name = parse_alarm_code(alarm['报警类型'])
            display_code = f"{code}：{name}"

            states_clean = parse_trigger_states(alarm.get('触发状态', ''))
            if len(states_clean) == 0 or (len(states_clean) == 1 and states_clean[0] == ''):
                states_clean = ['']

            for si, state_name in enumerate(states_clean):
                new_alarm_tbl = copy.deepcopy(tpl_alarm_tbl._tbl)
                body.append(new_alarm_tbl)
                alarm_tbl = doc.tables[-1]
                total_tables += 1

                set_cell_text_keep_valign(alarm_tbl.rows[0].cells[0], display_code, bold=True, size=12)

                trigger_method = alarm.get('报警原因及触发方法', '')
                test_condition_text = f'前置条件：{state_name}\n触发条件：{trigger_method}\n功能选择：\n操作：'
                set_cell_text_keep_valign(alarm_tbl.rows[2].cells[0], test_condition_text, size=9, align='left')

                alarm_text = alarm.get('报警文本', '')
                set_cell_text_keep_valign(alarm_tbl.rows[2].cells[2], alarm_text, size=9)

                if code == 'U23':
                    beep_text = '蜂鸣器短鸣10次'
                elif code == 'E7':
                    beep_text = '蜂鸣器长鸣'
                else:
                    beep_text = '蜂鸣器短鸣1次'
                set_cell_text_keep_valign(alarm_tbl.rows[3].cells[2], beep_text, size=9)

                set_cell_text_keep_valign(alarm_tbl.rows[4].cells[2], code, size=9)

                set_cell_text_keep_valign(alarm_tbl.rows[5].cells[2], '关闭电机', size=9)
                set_cell_text_keep_valign(alarm_tbl.rows[6].cells[2], '关闭加热', size=9)

                recoverable = alarm.get('是否可恢复', '')
                if '否' in str(recoverable):
                    handle_text = '重新上电'
                else:
                    if code == 'U2':
                        handle_text = '闭合杯盖'
                    elif code in ('E3', 'E4'):
                        handle_text = '电压调至220V'
                    elif code in ('E3A', 'E4A'):
                        handle_text = '软件屏蔽，电压调至220V'
                    elif code == 'U23':
                        handle_text = '正确放好接浆杯'
                    elif code == 'U24':
                        handle_text = '清理余水盒并放回'
                    elif code == 'U28':
                        handle_text = '水箱加水'
                    elif code == 'U39':
                        handle_text = '正确放置出浆嘴'
                    else:
                        handle_text = ''
                set_cell_text_keep_valign(alarm_tbl.rows[7].cells[2], handle_text, size=9)

                set_cell_text_keep_valign(alarm_tbl.rows[8].cells[2], recoverable, size=9)

                for row in alarm_tbl.rows:
                    for cell in row.cells:
                        paras_to_remove = []
                        for pi, para in enumerate(cell.paragraphs):
                            if pi == 0:
                                continue
                            if not para.text.strip():
                                paras_to_remove.append(para._p)
                        for p in paras_to_remove:
                            p.getparent().remove(p)

    existing_sectPr = body.find(qn('w:sectPr'))
    if existing_sectPr is not None:
        body.remove(existing_sectPr)
    if sectPr is not None:
        body.append(sectPr)

    doc.save(output_path)
    return total_tables


# ===================== Flask 路由 =====================

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/convert/menu', methods=['GET', 'POST'])
def convert_menu():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('请选择文件')
            return redirect(request.url)

        file = request.files['file']
        if file.filename == '':
            flash('请选择文件')
            return redirect(request.url)

        if not file.filename.endswith(('.xlsx', '.xls')):
            flash('仅支持 .xlsx 或 .xls 格式')
            return redirect(request.url)

        # 保存上传文件
        filename = file.filename
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)

        # 转换
        output_filename = filename.replace('.xlsx', '').replace('.xls', '') + '_一致性核对.xlsx'
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        try:
            sheet_count = convert_menu_consistency(input_path, output_path)
            flash(f'转换成功！共处理 {sheet_count} 个sheet。点击下方按钮下载结果。')
            return render_template('result.html',
                                download_url=url_for('download_file', filename=output_filename),
                                filename=output_filename)
        except Exception as e:
            flash(f'转换失败：{str(e)}')
            return redirect(request.url)

    return render_template('upload.html', title='菜单一致性核对转换', action=url_for('convert_menu'))

@app.route('/convert/alarm', methods=['GET', 'POST'])
def convert_alarm_route():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('请选择文件')
            return redirect(request.url)

        file = request.files['file']
        if file.filename == '':
            flash('请选择文件')
            return redirect(request.url)

        if not file.filename.endswith(('.xlsx', '.xls')):
            flash('仅支持 .xlsx 或 .xls 格式')
            return redirect(request.url)

        filename = file.filename
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)

        output_filename = filename.replace('.xlsx', '').replace('.xls', '') + '_报警用例.xlsx'
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        try:
            sheet_count = convert_alarm(input_path, output_path)
            flash(f'转换成功！共处理 {sheet_count} 个sheet。')
            return render_template('result.html',
                                download_url=url_for('download_file', filename=output_filename),
                                filename=output_filename)
        except Exception as e:
            flash(f'转换失败：{str(e)}')
            return redirect(request.url)

    return render_template('upload.html', title='报警用例转换（Excel）', action=url_for('convert_alarm_route'))


@app.route('/convert/alarm-word', methods=['GET', 'POST'])
def convert_alarm_word_route():
    """报警核对Word生成（完整版）"""
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('请选择文件')
            return redirect(request.url)

        file = request.files['file']
        if file.filename == '':
            flash('请选择文件')
            return redirect(request.url)

        if not file.filename.endswith('.docx'):
            flash('仅支持 .docx 格式（Word文档）')
            return redirect(request.url)

        filename = file.filename
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(input_path)

        template_path = '程序/程序/L15-Pxx-样机-报警核对-版本号 - 训练1.docx'

        # 自动版本号
        pattern = os.path.join(OUTPUT_FOLDER, '报警核对-完整报告-V*.docx')
        existing = glob.glob(pattern)
        max_ver = 0
        for f in existing:
            m = re.search(r'-V(\d+)\.docx', f)
            if m:
                v = int(m.group(1))
                if v > max_ver:
                    max_ver = v
        next_ver = max_ver + 1
        output_filename = f'报警核对-完整报告-V{next_ver}.docx'
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)

        # 标题信息
        title_info = {
            'model': request.form.get('model', 'DJ12-K6pro'),
            'version': request.form.get('version', ''),
            'checksum': request.form.get('checksum', ''),
            'voltage': request.form.get('voltage', '220V/50HZ'),
            'power': [
                request.form.get('power1', '加热功率：1200W'),
                request.form.get('power2', '搅拌功率：400W')
            ],
            'stage': request.form.get('stage', ''),
        }

        try:
            total = generate_full_alarm_report(template_path, input_path, output_path, title_info)
            flash(f'生成成功！共生成 {total} 个报警核对表。')
            return render_template('result.html',
                                download_url=url_for('download_file', filename=output_filename),
                                filename=output_filename)
        except Exception as e:
            flash(f'生成失败：{str(e)}')
            return redirect(request.url)

    return render_template('upload_alarm.html',
                          title='报警核对Word生成',
                          action=url_for('convert_alarm_word_route'))

@app.route('/download/<filename>')
def download_file(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    return send_file(path, as_attachment=True)

@app.route('/preview', methods=['POST'])
def preview():
    """预览转换结果（返回JSON）"""
    if 'file' not in request.files:
        return {'error': '请选择文件'}, 400

    file = request.files['file']
    convert_type = request.form.get('type', 'menu')

    if file.filename == '':
        return {'error': '请选择文件'}, 400

    filename = file.filename
    input_path = os.path.join(UPLOAD_FOLDER, '_preview_' + filename)
    file.save(input_path)

    try:
        wb = openpyxl.load_workbook(input_path)
        sheets_info = []
        for sheet_name in wb.sheetnames:
            if sheet_name == '目录':
                continue
            sheet = wb[sheet_name]
            rows = list(sheet.iter_rows(values_only=True))
            # 找到表头
            header = []
            data_preview = []
            for i, row in enumerate(rows):
                row_vals = [str(v)[:50] if v else '' for v in row]
                if not header and any('步骤' in v or '动作' in v for v in row_vals if v):
                    header = row_vals
                elif header and len(data_preview) < 5:
                    data_preview.append(row_vals)
                elif header and len(data_preview) >= 5:
                    break

            sheets_info.append({
                'name': sheet_name,
                'header': header,
                'preview': data_preview,
                'row_count': max(0, len(rows) - (1 if header else 0))
            })

        os.remove(input_path)
        return {'sheets': sheets_info, 'type': convert_type}
    except Exception as e:
        return {'error': str(e)}, 500


@app.route('/version')
def version():
    """版本验证接口"""
    return {
        'version': 'v4.0',
        'features': [
            '标题行居中',
            'D列下拉选择（符合/不符合）',
            '不符合时红底强调',
            '步骤序号自动补充',
            'B/C列合并支持'
        ],
        'status': 'latest'
    }

if __name__ == '__main__':
    # 生产环境使用环境变量 PORT，本地开发使用 5001
    import os
    port = int(os.environ.get('PORT', 5001))
    app.run(debug=False, port=port, host='0.0.0.0')  # 允许外部访问
