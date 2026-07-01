#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报警核对Word生成模块
从线路板功能说明书提取报警数据，生成报警核对Word文档
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def extract_alarm_data(source_docx_path):
    """
    从线路板功能说明书Word文档中提取报警数据
    :param source_docx_path: 源Word文档路径
    :return: 报警数据列表
    """
    doc = Document(source_docx_path)
    
    alarm_data = []
    
    # 找到报警表格（通常是第一个表格，包含"报警类型"列）
    for table in doc.tables:
        # 检查是否是报警表格
        if len(table.rows) > 1 and '报警类型' in str([cell.text for cell in table.rows[0].cells]):
            # 从第1行开始（第0行是表头）
            for row_idx in range(1, len(table.rows)):
                row = table.rows[row_idx]
                cells = [cell.text for cell in row.cells]
                
                if len(cells) >= 3:
                    # 提取报警类型（例如：U3\n(防溢报警)）
                    alarm_type = cells[1].strip()
                    
                    # 转换为格式：U3：防溢报警
                    match = re.match(r'([UE]\d+)\s*\(?\s*(.+?)\s*\)?', alarm_type)
                    if match:
                        alarm_code = f"{match.group(1)}：{match.group(2)}"
                    else:
                        alarm_code = alarm_type.replace('\n', '').replace('(', '：').replace(')', '')
                    
                    # 提取报警文本
                    alarm_text = cells[2].strip() if len(cells) > 2 else ''
                    
                    # 创建测试项（按照模板格式）
                    test_items = [
                        {'item': '报警文案', 'expected': '符合□ \n不符合□'},
                        {'item': '蜂鸣器报警', 'expected': '符合□ \n不符合□'},
                        {'item': '数码管显示内容', 'expected': '符合□ \n不符合□'},
                        {'item': '输出控制', 'expected': '符合□ \n不符合□'},
                        {'item': '输出控制', 'expected': '符合□ \n不符合□'},
                        {'item': '处理操作', 'expected': '符合□ \n不符合□'},
                        {'item': '报警恢复状态', 'expected': '符合□ \n不符合□'}
                    ]
                    
                    alarm_data.append({
                        'alarm_code': alarm_code,
                        'alarm_text': alarm_text,
                        'test_items': test_items
                    })
            break  # 只处理第一个报警表格
    
    return alarm_data


def set_cell_border(cell):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_elem = OxmlElement(f'w:{edge}')
        edge_elem.set(qn('w:val'), 'single')
        edge_elem.set(qn('w:sz'), '4')
        edge_elem.set(qn('w:color'), 'auto')
        tcBorders.append(edge_elem)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, bold=False, size=10):
    """设置单元格文本"""
    cell.text = text
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = bold
            run.font.size = size * 100  # pt to twips


def create_alarm_word(output_path, alarm_data, title_info=None):
    """
    创建报警核对Word文档
    :param output_path: 输出文件路径
    :param alarm_data: 报警数据列表
    :param title_info: 标题信息字典（可选）
    :return: 生成的报警数量
    """
    
    doc = Document()
    
    # ========== 表格0：标题信息 ==========
    table0 = doc.add_table(rows=3, cols=6)
    table0.style = 'Table Grid'
    
    # 行0：产品报警代码核查（合并A-F）
    cell_00 = table0.rows[0].cells[0]
    cell_00.text = '产品报警代码核查'
    # 合并A-F列
    for i in range(1, 6):
        cell_00.merge(table0.rows[0].cells[i])
    
    # 设置标题行样式
    set_cell_text(cell_00, '产品报警代码核查', bold=True, size=14)
    cell_00.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 行1：样机名称/型号、程序版本号、校验和
    set_cell_text(table0.rows[1].cells[0], '样机名称/型号', bold=True)
    set_cell_text(table0.rows[1].cells[2], '程序版本号', bold=True)
    set_cell_text(table0.rows[1].cells[4], '校验和', bold=True)
    
    # 如果提供了标题信息，填入
    if title_info:
        if 'model' in title_info:
            table0.rows[1].cells[1].text = title_info['model']
        if 'version' in title_info:
            table0.rows[1].cells[3].text = title_info['version']
        if 'checksum' in title_info:
            table0.rows[1].cells[5].text = title_info['checksum']
        if 'voltage' in title_info:
            table0.rows[2].cells[1].text = title_info['voltage']
        if 'power' in title_info:
            table0.rows[2].cells[3].text = title_info['power']
        if 'stage' in title_info:
            table0.rows[2].cells[5].text = title_info['stage']
    
    # 行2：额定电压、额定功率、阶段
    set_cell_text(table0.rows[2].cells[0], '额定电压', bold=True)
    set_cell_text(table0.rows[2].cells[2], '额定功率', bold=True)
    set_cell_text(table0.rows[2].cells[4], '阶段', bold=True)
    
    # 设置所有单元格居中
    for row in table0.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加空行
    doc.add_paragraph()
    
    # ========== 为每个报警创建一个表格 ==========
    for idx, alarm in enumerate(alarm_data):
        # 创建表格：9行 x 4列
        table = doc.add_table(rows=9, cols=4)
        table.style = 'Table Grid'
        
        # 行0：报警代码（合并A-D）
        cell_00 = table.rows[0].cells[0]
        cell_00.text = alarm['alarm_code']
        for i in range(1, 4):
            cell_00.merge(table.rows[0].cells[i])
        
        # 设置报警代码标题样式
        set_cell_text(cell_00, alarm['alarm_code'], bold=True, size=12)
        cell_00.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 行1：表头
        set_cell_text(table.rows[1].cells[0], '测试条件', bold=True)
        set_cell_text(table.rows[1].cells[1], '检核项', bold=True)
        set_cell_text(table.rows[1].cells[2], '预期结果', bold=True)
        set_cell_text(table.rows[1].cells[3], '实际结果', bold=True)
        
        # 行2-8：测试内容
        test_items = alarm.get('test_items', [
            {'item': '报警文案', 'expected': '符合□ \n不符合□'},
            {'item': '蜂鸣器报警', 'expected': '符合□ \n不符合□'},
            {'item': '数码管显示内容', 'expected': '符合□ \n不符合□'},
            {'item': '输出控制', 'expected': '符合□ \n不符合□'},
            {'item': '输出控制', 'expected': '符合□ \n不符合□'},
            {'item': '处理操作', 'expected': '符合□ \n不符合□'},
            {'item': '报警恢复状态', 'expected': '符合□ \n不符合□'}
        ])
        
        # 前置条件文本
        precond_text = '前置条件：\n\n触发条件：\n\n功能选择：\n\n操作:'
        
        # A列：合并行2-8
        cell_a = table.rows[2].cells[0]
        cell_a.text = precond_text
        for i in range(3, 9):  # 行3-8
            cell_a.merge(table.rows[i].cells[0])
        
        # 设置A列样式
        cell_a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell_a.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # B列、D列
        for i in range(7):
            row_idx = i + 2
            row = table.rows[row_idx]
            
            # B列：检核项
            if i < len(test_items):
                set_cell_text(row.cells[1], test_items[i]['item'])
            
            # C列：预期结果（空）
            # D列：实际结果
            if i < len(test_items):
                set_cell_text(row.cells[3], test_items[i]['expected'])
        
        # 设置所有单元格居中
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行（报警之间）
        if idx < len(alarm_data) - 1:
            doc.add_paragraph()
    
    # 保存文档
    doc.save(output_path)
    return len(alarm_data)
